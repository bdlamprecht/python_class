#!/usr/bin/env python

import os,pwd,time

from docx import Document
from docx.shared import Inches

document = Document()
recordset = {
    'test_name' : 'test_name',
    'device' :  'localhost',
    'tester' : "{} ({})".format(pwd.getpwuid(os.getuid())[4],
                                pwd.getpwuid(os.getuid())[0]),
    'date' : time.strftime("%d%b%Y"),}

document.add_heading('Document Title', 0)
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = 'row0'
    row_cells[1].text = 'row1'
    row_cells[2].text = 'row2'

document.add_page_break()

document.save('demo.docx')