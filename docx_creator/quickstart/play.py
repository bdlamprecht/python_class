#!/usr/bin/env python

import os,pwd,time
from docx import Document
from docx.shared import Inches

doc = Document('input.docx')

paragraph1 = doc.add_paragraph('This is a paragraph.')
prior = paragraph1.insert_paragraph_before('This is a prior paragraph.')

for i in range(0,2): {
	doc.add_heading('This is heading #{}.'.format(i+1))
	}

#doc.add_page_break()

#table = doc.add_table(rows=3,cols=2)
#table.style = 'LightShading-Accent1'

#cell1 = table.cell(0,0)
#cell1.text = 'parrot, possbily dead'

#cell2 = table.cell(0,1)
#cell2.text = 'brady is great'

#row = table.rows[1]
#row.cells[0].text = 'Foo bar to you.'
#row.cells[1].text = '...and to you, sir!'

doc.add_paragraph('{{ test_name }} - {{ device }} ', style='test_title')

paragraph2 = doc.add_paragraph('This is a test of the ')
run2 = paragraph2.add_run('emergency broadcast')
#run = paragraph2.add_run('emergency broadcast').underline = True
run2.bold = True
run2.italic = True
run2.underline = True
paragraph2.add_run(' system.')

paragraph3 = doc.add_paragraph('Normal text, ')
paragraph3.add_run('text with emphasis.', 'Emphasis')

paragraph4 = doc.add_paragraph('More normal tesxt, ')
run4 = paragraph4.add_run('more text with emphasis.')
run4.style = 'Emphasis'

doc.save('output.docx')