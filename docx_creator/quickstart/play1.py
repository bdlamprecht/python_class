#!/usr/bin/env python

import os,pwd,time
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt

doc = Document('input.docx')

#table = doc.add_table(rows=3,cols=2)
#table.style = 'LightShading-Accent1'

#cell1 = table.cell(0,0)
#cell1.text = 'parrot, possbily dead'

#cell2 = table.cell(0,1)
#cell2.text = 'brady is great'

#row = table.rows[1]
#row.cells[0].text = 'Foo bar to you.'
#row.cells[1].text = '...and to you, sir!'

# Demontrating basic paragraph functions

doc.add_paragraph('{{ test_name }} - {{ device }} ', style='test_title')

paragraph01 = doc.add_paragraph('This is a paragraph.')
prior = paragraph01.insert_paragraph_before('This is a prior paragraph.')

for i in range(0,2): {
	doc.add_heading('This is heading #{}.'.format(i+1))
	}

# Mixed text formatting.

paragraph02 = doc.add_paragraph('This is a test of the ')
run2 = paragraph02.add_run('emergency broadcast')
#run = paragraph02.add_run('emergency broadcast').underline = True
run2.bold = True
run2.italic = True
run2.underline = True
paragraph02.add_run(' system.')

# Additional text formatting.

paragraph03 = doc.add_paragraph('Normal text, ')
paragraph03.add_run('text with emphasis.', 'Emphasis')

paragraph04 = doc.add_paragraph('More normal text, ')
run4 = paragraph04.add_run('more text with emphasis.')
run4.style = 'Emphasis'

#doc.add_page_break()

# Text Alignment

paragraph05 = doc.add_paragraph("Left Alignment Text")
paragraph05_format = paragraph05.paragraph_format
paragraph05_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

paragraph06 = doc.add_paragraph("Center Alignment Text")
paragraph06_format = paragraph06.paragraph_format
paragraph06_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

paragraph07 = doc.add_paragraph("Right Alignment Text")
paragraph07_format = paragraph07.paragraph_format
paragraph07_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Text Indentation

paragraph08 = doc.add_paragraph("Left Indented Text")
paragraph08_format = paragraph08.paragraph_format
paragraph08_format.left_indent = Inches(0.5)

paragraph09 = doc.add_paragraph("Right Indented Text")
paragraph09_format = paragraph09.paragraph_format
paragraph09_format.right_indent = Pt(450)

# Complex Indentation

paragraph10 = doc.add_paragraph("This is a first line indented text.  Also demonstrating negative margins. A lot of text is needed to demontrate this phenomenon.")
paragraph10_format = paragraph10.paragraph_format
paragraph10_format.first_line_indent = Inches(-0.5)

# Tab Stops

paragraph11 = doc.add_paragraph("Showing creation of tab stops.")
paragraph11_format = paragraph11.paragraph_format
tab_stops11 = paragraph11_format.tab_stops
tab_stops11 = tab_stops11.add_tab_stop(Inches(1.5))
#print("Normal: {}".format(tab_stops11.position))
#print("Inches: {}".format(tab_stops11.position.inches))

# Showing alignment of tab stops - don't completely understand yet.

#paragraph12_format = paragraph12.paragraph_format
#tab_stops12 = paragraph12_format.tab_stops
#tab_stops12 = tab_stops12.add_tab_stop(Inches(1)), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS

# Showing usage of tab stops - also don't completely understand.

paragraph13 = doc.add_paragraph("Showing usage of tab stops.")
paragraph13_format = paragraph13.paragraph_format
tab_stops13 = paragraph13_format.tab_stops
tab_stops13 = tab_stops13.add_tab_stop(Inches(0.5))

doc.save('output1.docx')