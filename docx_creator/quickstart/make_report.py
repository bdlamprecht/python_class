#!/usr/bin/env python

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

doc = Document('input.docx')

styles = doc.styles

# Remove "latent" styles
#styles['Normal'].delete()
#styles['Heading 1'].delete()
#styles['Heading 2'].delete()
#styles['Heading 3'].delete()
#styles['HTML Preformatted'].delete()

paragraph_styles = [
  s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
]

table_styles = [
  s for s in styles if s.type == WD_STYLE_TYPE.TABLE
]

print('\nParapgrah Sytles:')
for style in paragraph_styles:
  print('Style Name: \t{}'.format(style.name))

print('\nTable Styles:')
for style in table_styles:
  print('Style Name: \t{}'.format(style.name))

for i in range(1,3):
  doc.add_paragraph('Test {}'.format(i), style='test_header')
  doc.add_paragraph('Description', style='desc_header')
  doc.add_paragraph('This is test {} text.'.format(i), style='normal_text')

  doc.add_paragraph('Commands To Execute', style='desc_header')
  doc.add_paragraph('Cisco:', style='bold')
  doc.add_paragraph('[ios command here]', style='output')
  doc.add_paragraph('Juniper:', style='bold')
  doc.add_paragraph('[junos command here]', style='output')

  doc.add_paragraph('Test Output', style='desc_header')

  # This section is to get the font size in the following table to be correct size.
  outp = doc.add_paragraph('This is some text')
  outp = doc.paragraphs[-1]
  outp.style.font.size = Pt(8)
  p = outp._element
  p.getparent().remove(p)
  p._p = p._element = None

  o_table = doc.add_table(1,1,style='test_output')
  o_table.cell(0,0).text = '{{ test{}_results }}'.format(i)
  #o_table.cell(0,0).add_paragraph('{{ test{}_results }}'.format(i),style='output_sm')

  # This section is to get the font size in the following table to be correct size.
  comm = doc.add_paragraph('This is some text')
  comm = doc.paragraphs[-1]
  comm.style.font.size = Pt(10)
  p = comm._element
  p.getparent().remove(p)
  p._p = p._element = None

  doc.add_paragraph('Test Comments', style='desc_header')
  c_table = doc.add_table(1,1,style='test_comment')
  c_table.cell(0,0).text = '{{ test{}_comment }}'.format(i)
  #c_table.cell(0,0).add_paragraph('{{ test{}_comment }}'.format(i),style='normal_text')

doc.save('report.docx')
