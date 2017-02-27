#!/usr/bin/env python

import os,pwd,time
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt, Length

doc = Document('input.docx')
doc.add_paragraph('{{ test_name }} - {{ device }} ', style='test_title')

# Showing paragraph spacing

paragraph01 = doc.add_paragraph("Showing paragraph spacing. 18pt before and 1.5 inches after.")
paragraph01_format = paragraph01.paragraph_format
paragraph01_format.space_before = Pt(18)
paragraph01_format.space_after = Inches(1.5)

# Showing line spacing

paragraph02 = doc.add_paragraph("Showing line spacing.")
paragraph02_format = paragraph02.paragraph_format
#print("Before")
#print("Line Spacing: \t\t{}".format(paragraph02_format.line_spacing))
#print("Line Spacing Rule: \t{}".format(paragraph02_format.line_spacing_rule))
paragraph02_format.line_spacing = Pt(18)
#print("1st After")
#print("Line Spacing: \t\t{}".format(paragraph02_format.line_spacing))
#print("Line Spacing Rule: \t{}".format(paragraph02_format.line_spacing_rule))
paragraph02_format.line_spacing = 1.75
#print("2nd After")
#print("Line Spacing: \t\t{}".format(paragraph02_format.line_spacing))
#print("Line Spacing Rule: \t{}".format(paragraph02_format.line_spacing_rule))

# Demonstrating pagination properties

paragraph03 = doc.add_paragraph("Showing pagination properties.")
paragraph03_format = paragraph03.paragraph_format
#print("Before")
#print("Keep Together: \t\t{}".format(paragraph03_format.keep_together))
#print("Keep With Next: \t{}".format(paragraph03_format.keep_with_next))
#print("Page Break Before: \t{}".format(paragraph03_format.page_break_before))
#print("Widow Control: \t\t{}".format(paragraph03_format.widow_control))
paragraph03_format.keep_together = True
paragraph03_format.keep_with_next = True
#paragraph03_format.page_break_before = True
paragraph03_format.widow_control = True
#print("After")
#print("Keep Together: \t\t{}".format(paragraph03_format.keep_together))
#print("Keep With Next: \t{}".format(paragraph03_format.keep_with_next))
#print("Page Break Before: \t{}".format(paragraph03_format.page_break_before))
#print("Widow Control: \t\t{}".format(paragraph03_format.widow_control))

doc.save('output2.docx')