#!/usr/bin/env python

import os,pwd,time
from docx import Document
from docx.shared import Inches, Pt, Length, RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.enum.section import WD_SECTION, WD_ORIENT

doc = Document('input.docx')
doc.add_paragraph('{{ test_name }} - {{ device }} ', style='test_title')

# Playing with sections

doc.add_paragraph('In the original section!')
sections = doc.sections
#print("Before")
#print("Sections: \t{}".format(len(sections)))
current01 = sections[-1]
#print(current01.start_type)
new01 = doc.add_section(WD_SECTION.ODD_PAGE)
doc.add_paragraph('In a new section!')
#print("After")
#print(new01.start_type)

# Playing with orientation - apparent bug with python-docx

new02 = doc.add_section()
pg_o = new02.orientation
pg_w = new02.page_width
pg_h = new02.page_height
print("Orientation: \t{}".format(pg_o))
print("Page Width: \t{}".format(pg_w))
print("Page Height: \t{}".format(pg_h))
new02.orientation = WD_ORIENT.LANDSCAPE
doc.add_paragraph('LANDSCAPE')
pg_o = new02.orientation
pg_w = new02.page_width
pg_h = new02.page_height
print("Orientation: \t{}".format(pg_o))
print("Page Width: \t{}".format(pg_w))
print("Page Height: \t{}".format(pg_h))
doc.add_section()

doc.save('output4.docx')