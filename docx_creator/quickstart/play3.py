#!/usr/bin/env python

import os,pwd,time
from docx import Document
from docx.shared import Inches, Pt, Length, RGBColor
from docx.enum.text import WD_UNDERLINE

doc = Document('input.docx')
doc.add_paragraph('{{ test_name }} - {{ device }} ', style='test_title')

# Showing character formatting

run01 = doc.add_paragraph().add_run("This is some text.")
font01 = run01.font
font01.name = 'Calibri'
font01.size = Pt(20)

# Showing additional character formatting

run02 = doc.add_paragraph().add_run("This is additional FORMATTED text.")
font02 = run02.font
font02.italic = True
font02.bold=True
font02.underline = True

# Showing underline formating options

run03 = doc.add_paragraph().add_run("This is wavy-heavy-underscore text.")
font03 = run03.font
font03.underline = WD_UNDERLINE.WAVY_HEAVY

# Showing font color options

run04 = doc.add_paragraph().add_run("This is colored text using hex-based values for RGBColor.")
font04 = run04.font
font04.color.rgb = RGBColor(0x42,0x24,0xE9)

run05 = doc.add_paragraph().add_run("This is colored text using decimal-based values for RGBColor.")
font05 = run05.font
font05.color.rgb = RGBColor(100,50,25)

doc.save('output3.docx')